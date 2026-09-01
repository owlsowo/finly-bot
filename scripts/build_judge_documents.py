#!/usr/bin/env python3
"""Build Finly's restrained, essay-first judge documents from Markdown."""

from __future__ import annotations

import re
import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ONE_PAGE_MD = ROOT / "docs/paper/one_page_writeup.md"
PAPER_MD = ROOT / "docs/paper/finly_engineering_appendix.md"
PUBLIC_DIR = ROOT / "public/judge"
DIST_DIR = ROOT / "dist/judge"

SERIF = "Times New Roman"
SANS = "Arial"
MONO = "Menlo"
MATH = "Cambria Math"
INK = "171717"
MUTED = "5F6368"
LINK = "1F5A4C"
RULE = "B8BBB8"
WASH = "F3F5F2"
GREEN = "246B57"
NAVY = "17324D"
DOCUMENT_TIMESTAMP = datetime(2026, 8, 30, 12, 0, 0, tzinfo=timezone.utc)

INLINE = re.compile(r"(\[[^\]]+\]\([^)]+\)|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def set_run_font(run, name: str, size: float, color: str = INK, *, bold=False, italic=False) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph_rule(paragraph, *, color: str = RULE, size: int = 4, space: int = 2) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, color: str = WASH) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, *, top=72, start=80, bottom=72, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_table_row_together(row, *, repeat_header=False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    if repeat_header:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def add_hyperlink(paragraph, label: str, url: str, *, size: float, font: str, bold=False) -> None:
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), font)
    run_fonts.set(qn("w:hAnsi"), font)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK)
    run_size = OxmlElement("w:sz")
    run_size.set(qn("w:val"), str(round(size * 2)))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([run_fonts, color, run_size, underline])
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, *, size: float, font: str = SERIF, color: str = INK, bold=False) -> None:
    cursor = 0
    for match in INLINE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, font, size, color, bold=bold)
        token = match.group(0)
        if token.startswith("["):
            parsed = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if parsed is None:
                run = paragraph.add_run(token)
                set_run_font(run, font, size, color, bold=bold)
            else:
                label, url = parsed.groups()
                add_hyperlink(paragraph, label, url, size=size, font=font, bold=bold)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, MONO, max(7.0, size - 0.6), color, bold=bold)
        elif token.startswith("**"):
            add_inline(paragraph, token[2:-2], size=size, font=font, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, font, size, color, bold=bold, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, font, size, color, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])
    set_run_font(run, SERIF, 8.0, MUTED)


def set_letter_page(section, *, top: float, bottom: float, left: float, right: float) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)


def configure_normal_style(doc: Document, *, size: float, line_spacing: float, after: float) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), SERIF)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), SERIF)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), SERIF)
    normal.font.size = Pt(size)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(after)
    normal.paragraph_format.line_spacing = line_spacing
    normal.paragraph_format.widow_control = True


def add_quiet_footer(section, *, first_page=False) -> None:
    footer = section.first_page_footer if first_page else section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_page_number(paragraph)


def add_running_header(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run("FINLY  /  AI TRADING WITH BUILT-IN CHECKS")
    set_run_font(run, SANS, 7.4, MUTED)
    set_paragraph_rule(paragraph, color=RULE, size=3, space=2)


def split_blocks(lines: list[str], start: int) -> list[str]:
    blocks: list[str] = []
    buffer: list[str] = []
    for raw in lines[start:]:
        line = raw.strip()
        if line.startswith("## "):
            if buffer:
                blocks.append(" ".join(buffer))
                buffer.clear()
            blocks.append(line)
        elif not line:
            if buffer:
                blocks.append(" ".join(buffer))
                buffer.clear()
        else:
            buffer.append(line)
    if buffer:
        blocks.append(" ".join(buffer))
    return blocks


def set_metadata(doc: Document, *, title: str, subject: str) -> None:
    properties = doc.core_properties
    properties.title = title
    properties.author = "Bruce Wen"
    properties.last_modified_by = "Bruce Wen"
    properties.subject = subject
    properties.comments = "Prepared by Bruce Wen for the Alpaca AI Trading Agents Hackathon."
    properties.keywords = "Finly, Alpaca, AI trading, options, risk controls"
    properties.created = DOCUMENT_TIMESTAMP
    properties.modified = DOCUMENT_TIMESTAMP


def build_one_page() -> Path:
    lines = ONE_PAGE_MD.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    byline = lines[2].replace("[bwen412@brandeis.edu](mailto:bwen412@brandeis.edu)", "bwen412@brandeis.edu")
    scope = lines[4]
    blocks = split_blocks(lines, 6)

    doc = Document()
    section = doc.sections[0]
    set_letter_page(section, top=0.68, bottom=0.62, left=0.88, right=0.88)
    configure_normal_style(doc, size=11.0, line_spacing=1.15, after=7.0)
    add_quiet_footer(section)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, SERIF, 18.5, INK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1.5)
    run = paragraph.add_run(byline)
    set_run_font(run, SERIF, 9.2, INK)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(scope)
    set_run_font(run, SERIF, 8.4, MUTED, italic=True)
    set_paragraph_rule(paragraph, color=RULE, size=3, space=4)

    in_references = False
    first_body = True
    for block in blocks:
        if block == "## References":
            in_references = True
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run("References")
            set_run_font(run, SERIF, 8.8, INK, bold=True)
            continue

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.widow_control = True
        paragraph.paragraph_format.keep_together = not in_references
        if in_references:
            paragraph.paragraph_format.left_indent = Inches(0.14)
            paragraph.paragraph_format.first_line_indent = Inches(-0.14)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, block, size=7.5, color=MUTED)
        else:
            paragraph.paragraph_format.first_line_indent = Inches(0 if first_body else 0.21)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(7.0)
            add_inline(paragraph, block, size=11.0)
            first_body = False

    set_metadata(doc, title=title, subject="One-page narrative proposal")
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    DIST_DIR.mkdir(parents=True, exist_ok=True)
    output = PUBLIC_DIR / "Finly_Judge_Proposal.docx"
    doc.save(output)
    shutil.copy2(output, DIST_DIR / output.name)
    return output


def markdown_blocks(lines: list[str], start: int) -> list[tuple[str, object]]:
    blocks: list[tuple[str, str]] = []
    buffer: list[str] = []
    code_buffer: list[str] = []
    equation_buffer: list[str] = []
    table_buffer: list[str] = []
    in_code = False
    in_equation = False

    def flush() -> None:
        if buffer:
            blocks.append(("paragraph", " ".join(buffer)))
            buffer.clear()

    def flush_table() -> None:
        if not table_buffer:
            return
        parsed: list[list[str]] = []
        for row in table_buffer:
            cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                continue
            parsed.append(cells)
        if parsed:
            blocks.append(("table", parsed))
        table_buffer.clear()

    for raw in lines[start:]:
        line = raw.strip()
        if line.startswith("```") or line.startswith("~~~"):
            flush()
            flush_table()
            if in_code:
                blocks.append(("code", "\n".join(code_buffer)))
                code_buffer.clear()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buffer.append(raw.rstrip())
            continue
        if line == "$$":
            flush()
            flush_table()
            if in_equation:
                blocks.append(("equation", " ".join(equation_buffer)))
                equation_buffer.clear()
                in_equation = False
            else:
                in_equation = True
            continue
        if in_equation:
            equation_buffer.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            flush()
            table_buffer.append(line)
            continue
        flush_table()
        if line.startswith("### "):
            flush()
            blocks.append(("h2", line[4:].strip()))
        elif line.startswith("## "):
            flush()
            blocks.append(("h1", line[3:].strip()))
        elif line.startswith("- "):
            flush()
            blocks.append(("bullet", line[2:].strip()))
        elif re.match(r"^\d+\.\s+", line):
            flush()
            match = re.match(r"^(\d+)\.\s+(.+)", line)
            blocks.append(("number", (match.group(1), match.group(2))))
        elif not line:
            flush()
        else:
            buffer.append(line)
    flush()
    flush_table()
    if code_buffer:
        blocks.append(("code", "\n".join(code_buffer)))
    if equation_buffer:
        blocks.append(("equation", " ".join(equation_buffer)))
    return blocks


def add_technical_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        keep_table_row_together(table.rows[row_index], repeat_header=row_index == 0)
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, WASH)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            value = values[column_index] if column_index < len(values) else ""
            add_inline(
                paragraph,
                value,
                size=7.9 if row_index else 8.0,
                font=SANS,
                color="FFFFFF" if row_index == 0 else INK,
                bold=row_index == 0,
            )
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(0)


def add_section_heading(doc: Document, text: str, *, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(13)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(text)
        set_run_font(run, SERIF, 13.5, INK, bold=True)
    else:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        set_run_font(run, SERIF, 11.4, INK, bold=True, italic=True)


def build_technical_paper() -> Path:
    lines = PAPER_MD.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    subtitle = lines[2].removeprefix("## ").strip()
    byline = lines[4].replace("[bwen412@brandeis.edu](mailto:bwen412@brandeis.edu)", "bwen412@brandeis.edu")
    date_line = lines[6]
    blocks = markdown_blocks(lines, 8)

    doc = Document()
    section = doc.sections[0]
    set_letter_page(section, top=0.92, bottom=0.86, left=1.0, right=1.0)
    configure_normal_style(doc, size=11.0, line_spacing=1.15, after=6.0)
    section.different_first_page_header_footer = True
    add_running_header(section)
    add_quiet_footer(section)
    add_quiet_footer(section, first_page=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_run_font(run, SERIF, 23.5, INK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(9)
    run = paragraph.add_run(subtitle)
    set_run_font(run, SERIF, 11.8, MUTED, italic=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(byline)
    set_run_font(run, SERIF, 9.1, INK)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(date_line)
    set_run_font(run, SERIF, 8.3, MUTED)
    set_paragraph_rule(paragraph, color=RULE, size=3, space=5)

    in_abstract = False
    in_references = False
    first_after_heading = True
    for kind, value in blocks:
        if kind == "h1":
            add_section_heading(doc, value, level=1)
            in_abstract = value == "Abstract"
            in_references = value == "References"
            first_after_heading = True
            continue
        if kind == "h2":
            add_section_heading(doc, value, level=2)
            first_after_heading = True
            continue

        if kind == "table":
            add_technical_table(doc, value)
            first_after_heading = False
            continue

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.widow_control = True
        paragraph.paragraph_format.keep_together = not in_references
        if kind == "equation":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.right_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.keep_together = True
            set_paragraph_shading(paragraph, "F7F8F6")
            run = paragraph.add_run(str(value))
            set_run_font(run, MATH, 10.2, INK)
        elif kind == "code":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.right_indent = Inches(0.22)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.keep_together = True
            set_paragraph_shading(paragraph, WASH)
            set_paragraph_rule(paragraph, color=GREEN, size=6, space=4)
            for index, line in enumerate(str(value).splitlines()):
                run = paragraph.add_run(line)
                set_run_font(run, MONO, 7.7, INK)
                if index < len(str(value).splitlines()) - 1:
                    run.add_break()
        elif kind in {"bullet", "number"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.38)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.space_after = Pt(3.5)
            if kind == "bullet":
                add_inline(paragraph, "• " + str(value), size=10.35)
            else:
                number, body = value
                add_inline(paragraph, f"{number}. " + str(body), size=10.35)
        elif in_references:
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(-0.24)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(2.1)
            add_inline(paragraph, value, size=8.65, color=MUTED)
        elif in_abstract:
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.18)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.space_after = Pt(5)
            add_inline(paragraph, value, size=10.25)
        else:
            paragraph.paragraph_format.first_line_indent = Inches(0 if first_after_heading else 0.24)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6.0)
            add_inline(paragraph, value, size=11.0)
        first_after_heading = False

    set_metadata(doc, title=title, subject="Engineering appendix on controlled delegation and quantitative evidence")
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    DIST_DIR.mkdir(parents=True, exist_ok=True)
    output = PUBLIC_DIR / "Finly_Engineering_Appendix.docx"
    doc.save(output)
    shutil.copy2(output, DIST_DIR / output.name)
    return output


if __name__ == "__main__":
    print(build_one_page())
    print(build_technical_paper())
